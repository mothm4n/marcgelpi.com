#!/usr/bin/env python3

"""Generate the tagged PDF for the How to sell OKRs internally field guide."""

from __future__ import annotations

import argparse
import json
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


CHARCOAL = "252522"
PETROL = "174C5B"
COPPER = "C96C4D"
MUTED = "6B6A64"
PAPER = "F4EFE6"
WHITE = "FFFFFF"


def set_paragraph_border(paragraph, color: str = "CFC8BC") -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_run_color(run, color: str) -> None:
    run.font.color.rgb = RGBColor.from_string(color)


def set_language(style, language: str = "en-US") -> None:
    properties = style.element.get_or_add_rPr()
    language_element = properties.find(qn("w:lang"))
    if language_element is None:
        language_element = OxmlElement("w:lang")
        properties.append(language_element)
    language_element.set(qn("w:val"), language)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), PETROL)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([properties, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_image(document: Document, path: Path, width: float, alt: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inline_shape = paragraph.add_run().add_picture(str(path), width=Inches(width))
    inline_shape._inline.docPr.set("descr", alt)


def add_label(document: Document, text: str) -> None:
    document.add_paragraph(text.upper(), style="Resource label")


def add_quote(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Resource quote")
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), PETROL)
    properties.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "64")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), PETROL)
        borders.append(border)
    properties.append(borders)
    run = paragraph.add_run(text)
    set_run_color(run, WHITE)


def add_answer_lines(document: Document, count: int = 2) -> None:
    for _ in range(count):
        paragraph = document.add_paragraph(" ")
        paragraph.paragraph_format.space_after = Pt(7)
        set_paragraph_border(paragraph)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(CHARCOAL)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15
    set_language(normal)

    title = document.styles["Title"]
    title.font.name = "Georgia"
    title.font.size = Pt(32)
    title.font.bold = False
    title.font.color.rgb = RGBColor.from_string(CHARCOAL)
    title.paragraph_format.space_after = Pt(12)
    set_language(title)

    heading_one = document.styles["Heading 1"]
    heading_one.font.name = "Georgia"
    heading_one.font.size = Pt(27)
    heading_one.font.bold = False
    heading_one.font.color.rgb = RGBColor.from_string(CHARCOAL)
    heading_one.paragraph_format.space_before = Pt(0)
    heading_one.paragraph_format.space_after = Pt(14)
    set_language(heading_one)

    heading_two = document.styles["Heading 2"]
    heading_two.font.name = "Arial"
    heading_two.font.size = Pt(12)
    heading_two.font.bold = True
    heading_two.font.color.rgb = RGBColor.from_string(PETROL)
    heading_two.paragraph_format.space_before = Pt(12)
    heading_two.paragraph_format.space_after = Pt(4)
    set_language(heading_two)

    label = document.styles.add_style("Resource label", WD_STYLE_TYPE.PARAGRAPH)
    label.font.name = "Arial"
    label.font.size = Pt(8)
    label.font.bold = True
    label.font.color.rgb = RGBColor.from_string(COPPER)
    label.paragraph_format.space_after = Pt(7)
    set_language(label)

    quote = document.styles.add_style("Resource quote", WD_STYLE_TYPE.PARAGRAPH)
    quote.font.name = "Georgia"
    quote.font.size = Pt(17)
    quote.font.color.rgb = RGBColor.from_string(WHITE)
    quote.paragraph_format.left_indent = Cm(0)
    quote.paragraph_format.right_indent = Cm(0)
    quote.paragraph_format.space_before = Pt(0)
    quote.paragraph_format.space_after = Pt(0)
    quote.paragraph_format.line_spacing = 1.15
    set_language(quote)

    header = section.header.paragraphs[0]
    header.text = "HOW TO SELL OKRS INTERNALLY"
    header.style = normal
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    set_run_color(header.runs[0], PETROL)

    footer = section.footer.paragraphs[0]
    footer.add_run("Marc Gelpi  -  marcgelpi.com")
    footer.runs[0].font.size = Pt(8)
    set_run_color(footer.runs[0], MUTED)
    footer.add_run("                                      ")
    add_page_number(footer)


def add_cover(document: Document, hero_path: Path) -> None:
    add_image(
        document,
        hero_path,
        7.0,
        "Abstract petrol, charcoal and copper forms converging around a precise area of negative space.",
    )
    add_label(document, "Field guide  /  15 minutes")
    document.add_paragraph("How to sell OKRs internally", style="Title")
    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        "Build the case for focus, alignment, accountability and ambitious learning - without selling OKRs as a cure-all."
    )
    run.font.name = "Georgia"
    run.font.size = Pt(16)
    document.add_paragraph("Marc Gelpi  -  marcgelpi.com")
    document.add_page_break()


def add_case_page(document: Document, resource: dict) -> None:
    document.add_heading("Sell the outcome, not the acronym", level=1)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        "Leaders do not need another framework. They need fewer competing priorities, clearer connections between teams and earlier evidence that an important bet is drifting."
    )
    run.font.name = "Georgia"
    run.font.size = Pt(15)
    add_quote(document, "Do not start with cadence, templates or terminology. Start with the cost of the current system.")
    add_label(document, "Observable costs")
    add_bullets(document, resource["observable_costs"][:5])
    paragraph = document.add_paragraph()
    run = paragraph.add_run(
        "Choose two or three signals you can show. Evidence makes a stronger case than saying 'we need more alignment.'"
    )
    run.bold = True
    set_run_color(run, PETROL)
    document.add_page_break()


def add_outcomes_page(document: Document, resource: dict) -> None:
    document.add_heading("Four outcomes worth buying", level=1)
    for outcome in resource["outcomes"]:
        document.add_heading(f"{outcome['number']}  {outcome['title']}", level=2)
        document.add_paragraph(outcome["summary"])
    add_quote(document, "Which of these outcomes solves a problem your leaders already recognize? Start there.")
    document.add_page_break()


def add_diagnosis_page(document: Document) -> None:
    document.add_heading("Diagnose before you pitch", level=1)
    document.add_paragraph(
        "Describe the problem as observable behavior. Do not diagnose the organization with a framework it has not chosen."
    )
    prompts = [
        "What recurring priority or coordination problem can you show?",
        "What does that problem cost in time, trust, quality or opportunity?",
        "Which behavior should change if an OKR pilot is useful?",
        "What evidence would show that the behavior improved?",
    ]
    for index, prompt in enumerate(prompts, start=1):
        document.add_heading(f"{index}. {prompt}", level=2)
        add_answer_lines(document)
    document.add_page_break()


def add_conversation_page(document: Document, resource: dict) -> None:
    document.add_heading("Build the case in five moves", level=1)
    for move in resource["conversation_moves"]:
        document.add_heading(f"{move['number']}  {move['title']}", level=2)
        document.add_paragraph(move["prompt"])
        set_paragraph_border(document.paragraphs[-1])
    add_quote(document, "Frame OKRs as a test of better management behavior, not as a fashionable destination.")
    document.add_page_break()


def add_objections_page(document: Document, resource: dict) -> None:
    document.add_heading("Prepare for the obvious objections", level=1)
    for objection in resource["objections"]:
        document.add_heading(objection["title"], level=2)
        document.add_paragraph(objection["response"])
        set_paragraph_border(document.paragraphs[-1])
    document.add_heading("What is the strongest objection in your organization?", level=2)
    add_answer_lines(document, 3)
    document.add_page_break()


def add_pilot_page(document: Document, resource: dict) -> None:
    document.add_heading("Make a smaller ask", level=1)
    document.add_paragraph(
        "Do not ask the whole company to adopt a belief. Ask for one cycle, one strategic area and a decision based on evidence."
    )
    for field in ["Strategic area", "Pilot owner", "Cycle and review date", "Existing process this pilot replaces or simplifies"]:
        document.add_heading(field, level=2)
        add_answer_lines(document, 2)
    add_label(document, "Success signals")
    add_bullets(document, resource["success_signals"])
    document.add_page_break()


def add_pitch_page(document: Document, resource: dict) -> None:
    document.add_heading("A pitch you can use", level=1)
    add_quote(document, resource["pitch"])
    document.add_heading("Your decision request", level=2)
    add_answer_lines(document, 4)
    document.add_heading("Further reading", level=2)
    source = resource["source"]
    paragraph = document.add_paragraph(f"{source['attribution']}, ")
    add_hyperlink(paragraph, source["title"], source["url"])
    document.add_paragraph(source["guardrail"])


def build_document(resource: dict, hero_path: Path, destination: Path) -> None:
    document = Document()
    configure_document(document)
    properties = document.core_properties
    properties.title = "How to sell OKRs internally"
    properties.author = "Marc Gelpi"
    properties.subject = "A practical field guide for building internal support for OKRs"
    properties.keywords = "OKRs, leadership, alignment, focus, accountability, change"
    properties.language = "en-US"
    add_cover(document, hero_path)
    add_case_page(document, resource)
    add_outcomes_page(document, resource)
    add_diagnosis_page(document)
    add_conversation_page(document, resource)
    add_objections_page(document, resource)
    add_pilot_page(document, resource)
    add_pitch_page(document, resource)
    document.save(destination)


def export_tagged_pdf(source: Path, output: Path) -> None:
    soffice = shutil.which("soffice")
    if not soffice:
        raise RuntimeError("LibreOffice is required to export a tagged PDF")
    filter_options = (
        'pdf:writer_pdf_Export:{"UseTaggedPDF":{"type":"boolean","value":"true"},'
        '"PDFUACompliance":{"type":"boolean","value":"true"}}'
    )
    subprocess.run(
        [
            soffice,
            "--headless",
            f"-env:UserInstallation={source.parent.joinpath('libreoffice-profile').as_uri()}",
            "--convert-to",
            filter_options,
            "--outdir",
            str(source.parent),
            str(source),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    generated = source.with_suffix(".pdf")
    if not generated.exists():
        raise RuntimeError("LibreOffice did not generate the expected PDF")
    output.parent.mkdir(parents=True, exist_ok=True)
    generated.replace(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("output", nargs="?", default="output/pdf/how-to-sell-okrs.pdf")
    args = parser.parse_args()
    repo_root = Path(__file__).resolve().parent.parent
    resource = json.loads((repo_root / "data/resources/how_to_sell_okrs.json").read_text())
    temporary_root = repo_root / "tmp/pdfs"
    temporary_root.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="okrs-resource-", dir=temporary_root) as temporary_directory:
        source = Path(temporary_directory) / "how-to-sell-okrs.docx"
        build_document(resource, repo_root / "assets/images/resources/okrs-focus-abstract.png", source)
        export_tagged_pdf(source, repo_root / args.output)


if __name__ == "__main__":
    main()
